# extractors/dictionary.py
import json
import os
import re
from typing import Dict, List, Optional
from difflib import SequenceMatcher

try:
    import docx  # python-docx
except Exception:
    docx = None

DICT_JSON_PATH = os.getenv("DICT_JSON_PATH", "data/actors.json")

def _similar(a: str, b: str) -> float:
    return SequenceMatcher(None, a.upper(), b.upper()).ratio()

def load_dictionary_from_docx(docx_path: str) -> List[Dict]:
    """
    Lê DOCX com mapeamento CLIENTE x CÉLULA x PARTE INTERESSADA.
    Espera um conteúdo tabular ou linhas 'Cliente: ... / Célula: ... / Parte: ...'
    """
    items = []
    if not docx or not os.path.exists(docx_path):
        return items

    doc = docx.Document(docx_path)
    # 1) Tabelas
    for table in doc.tables:
        headers = [c.text.strip() for c in table.rows[0].cells]
        for row in table.rows[1:]:
            values = [c.text.strip() for c in row.cells]
            rowdict = dict(zip(headers, values))
            cliente = rowdict.get("CLIENTE") or rowdict.get("Cliente") or ""
            celula = rowdict.get("CÉLULA") or rowdict.get("Celula") or rowdict.get("Célula") or ""
            parte  = rowdict.get("PARTE INTERESSADA") or rowdict.get("Parte Interessada") or ""
            if any([cliente, celula, parte]):
                items.append({"cliente": cliente, "celula": celula, "parte_interessada": parte})

    # 2) Linhas soltas (fallback)
    buf = "\n".join(p.text for p in doc.paragraphs)
    for line in buf.splitlines():
        m = re.search(r"Cliente\s*:\s*(.+?)\s*/\s*C[ÉE]lula\s*:\s*(.+?)\s*/\s*Parte\s*Interessada\s*:\s*(.+)$", line, re.I)
        if m:
            items.append({
                "cliente": m.group(1).strip(),
                "celula":  m.group(2).strip(),
                "parte_interessada": m.group(3).strip(),
            })
    return items

def save_dictionary(items: List[Dict]):
    os.makedirs(os.path.dirname(DICT_JSON_PATH) or ".", exist_ok=True)
    with open(DICT_JSON_PATH, "w", encoding="utf-8") as f:
        json.dump(items, f, ensure_ascii=False, indent=2)

def load_dictionary() -> List[Dict]:
    if not os.path.exists(DICT_JSON_PATH):
        return []
    try:
        with open(DICT_JSON_PATH, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return []

def enrich_with_dictionary(data: Dict, dict_items: Optional[List[Dict]] = None) -> Dict:
    """
    Se parte_interessada bater com o dicionário, preenche cliente/célula.
    """
    if dict_items is None:
        dict_items = load_dictionary()

    parte = (data.get("parte_interessada") or "").strip()
    if not parte or not dict_items:
        return data

    best, bestscore = None, 0.0
    for item in dict_items:
        pi = (item.get("parte_interessada") or "").strip()
        if not pi:
            continue
        score = _similar(parte, pi)
        if score > bestscore:
            best, bestscore = item, score

    if best and bestscore >= 0.90:
        data.setdefault("cliente", best.get("cliente", ""))
        data.setdefault("celula", best.get("celula", ""))

    return data
